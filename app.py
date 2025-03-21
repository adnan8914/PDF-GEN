import streamlit as st
from docx import Document
from datetime import datetime
import os
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import uuid
import tempfile

# Proposal configurations
PROPOSAL_CONFIG = {
    "Make, Manychat & CRM Automation": {
        "template": "Make, Manychat & CRM Automation.docx",
        "pricing_fields": [
            ("ManyChat Automation", "MC-Price"),
            ("Make Automation", "M-Price"),
            ("CRM Automations", "C-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "Make & Manychat Automation": {
        "template": "Make & Manychat Automation.docx",
        "pricing_fields": [
            ("ManyChat Automation", "MC-Price"),
            ("Make Automation", "M-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "Ai Calling, Make, Manychat and CRM Automation": {
        "template": "Ai Calling, Make, Manychat and CRM Automation.docx",
        "pricing_fields": [
            ("AI Calling + CRM Integration", "AI-Price"),
            ("ManyChat & Make Automation", "MM-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "AI Calling, Make & CRM Automation": {
        "template": "AI Calling, Make & CRM Automation.docx",
        "pricing_fields": [
            ("AI Calling", "AI-Price"),
            ("Make Automation", "M-Price"),
            ("CRM Automations", "C-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "AI Calling(Basic) & CRM Automation": {
        "template": "AI Calling(Basic) & CRM Automation.docx",
        "pricing_fields": [
            ("AI Calling(Basic)", "AI-Price"),
            ("CRM Automation", "CC-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "AI Calling, Make & Manychat Automation": {
        "template": "Ai Calling, Make & Manychat Automation.docx",
        "pricing_fields": [
            ("AI Calling(Basic)", "AI-Price"),
            ("ManyChat Automation", "MC-Price"),
            ("Make Automation", "M-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "Ai Calling + CRM Intergration, Make & Manychat Automation, CRM Automation": {
        "template": "Ai Calling + CRM Intergration, Make & Manychat Automation, CRM Automation.docx",
        "pricing_fields": [
            ("AI Calling + CRM Integration", "AI-Price"),
            ("ManyChat & Make Automation", "MM-Price"),
            ("CRM Automation", "CC-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "AI Calling(Basic) & CRM Automation & Email Automation": {
        "template": "AI Calling(Basic) & CRM Automation & Email Automation.docx",
        "pricing_fields": [
            ("AI Calling(Basic)", "AI-Price"),
            ("CRM Automation", "CC-Price"),
            ("Email Automation", "E-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "Manychat & CRM Automation": {
        "template": "Manychat & CRM Automation.docx",
        "pricing_fields": [
            ("ManyChat Automation", "MC-Price"),
            ("CRM Automations", "C-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "Make & CRM Automation": {
        "template": "Make & CRM Automation.docx",
        "pricing_fields": [
            ("Make Automation", "M-Price"),
            ("CRM Automations", "C-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "Make, CRM Automation & AI Content Creation": {
        "template": "Make, CRM Automation & AI Content Creation.docx",
        "pricing_fields": [
            ("Make Automation", "M-Price"),
            ("CRM Automations", "C-Price"),
            ("AI Content Creation", "ACC-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "Digital Marketing": {
        "template": "DM Proposal.docx",
        "pricing_fields": [
            ("Marketing Strategy", "ms_price"),
            ("Social Media Handles Setup", "smh_price"),
            ("Meta & Google Ads Manager Setup", "sgam_price"),
            ("Creative Posts (10 per month)", "cp_price"),
            ("Meta Paid Ads", "ma_price"),
            ("Google Paid Ads", "gpa_price"),
            ("SEO", "seo_price"),
            ("Email Marketing", "em_price"),
            ("Monthly Maintenance & Reporting", "mmr_price")
        ],
        "team_fields": [
            ("Digital Marketing Executive", "dm_ex_no"),
            ("Digital Marketing Associate", "dm_asso_no"),
            ("Business Analyst", "ba_no"),
            ("Graphics Designer", "gd_no")
        ],
        "special_fields": [
            ("validity_date", "<<"),
            ("advnc_pay", "<<"),
            ("balnc_pay", "<<")
        ],
        "team_type": "digital_marketing"
    },
    "Shopify Website": {
        "template": "Shopify website.docx",
        "pricing_fields": [
            ("Development", "development"),
            ("Design", "design"),
            ("Testing & Live", "testing"),
            ("Annual Maintenance", "annual_mai"),
            ("Additional Features & Enhancements", "add_feature")
        ],
        "team_fields": [
            ("Project Manager", "pm_no"),
            ("Business Analyst", "ba_no"),
            ("UI/UX Members", "uix_no"),
            ("Backend Developers", "bd_no"),
            ("Frontend Developers", "fd_no"),
            ("AI/ML Developers", "aiml_no"),
            ("System Architect", "sa_no"),
            ("Shopify Developers", "sd_no")
        ],
        "special_fields": [
            ("validity_date", "<<"),
            ("location", "<<")  # Added for country/location field
        ],
        "team_type": "shopify"  # New team type for Shopify projects
    },
    "Web Based AI Fintech": {
        "template": "Web based AI Fintech proposal.docx",
        "pricing_fields": [
            ("Design", "design"),
            ("Development", "development"),
            ("AI/ML Models", "ai_ml_model"),
            ("Additional Features & Enhancements", "additional_feat")
        ],
        "team_fields": [
            ("Project Manager", "pm_no"),
            ("Business Analyst", "ba_no"),
            ("UI/UX Members", "uix_no"),
            ("Backend Developers", "bd_no"),
            ("Frontend Developers", "fd_no"),
            ("AI/ML Developers", "aiml_no"),
            ("System Architect", "sa_no"),
            ("AWS Developer", "aws_no")
        ],
        "special_fields": [
            ("validity_date", "<<"),
            ("client_location", "<<")  # Added for location field
        ],
        "team_type": "fintech"  # New team type for Fintech projects
    },
    "AI Based Search Engine": {
        "template": "AI Based Search Engine Website Technical Consultation proposal.docx",
        "pricing_fields": [
            ("Designs", "design"),
            ("Development", "development"),
            ("Testing & Deployment", "test_deploy"),
            ("Additional Features & Enhancements", "ad_f")
        ],
        "team_fields": [
            ("Project Manager", "pm_no"),
            ("Business Analyst", "ba_no"),
            ("UI/UX Members", "uix_no"),
            ("Backend Developers", "bd_no"),
            ("Frontend Developers", "fd_no"),
            ("AI/ML Developers", "aiml_no"),
            ("System Architect", "sa_no"),
            ("AWS Developer", "aws_no")
        ],
        "special_fields": [
            ("validity_date", "<<"),
            ("client_location", "<<")
        ],
        "team_type": "search_engine"
    },
    "Single Vendor Ecommerce": {
        "template": "Single Vendor Ecommerce website.docx",
        "pricing_fields": [
            ("Design", "design"),
            ("Development", "dev"),
            ("Website Chatbot", "wb_cb"),
            ("Testing & Deployment", "testing"),
            ("Additional Features & Enhancements", "ad_fs")
        ],
        "team_fields": [
            ("Project Manager", "pm_no"),
            ("Business Analyst", "ba_no"),
            ("UI/UX Members", "uix_no"),
            ("Backend Developers", "bd_no"),
            ("Frontend Developers", "fd_no"),
            ("AI/ML Developers", "aiml_no"),
            ("System Architect", "sa_no"),
            ("AWS Developer", "aws_no")
        ],
        "special_fields": [
            ("validity_date", "<<"),
            ("location", "<<")
        ],
        "team_type": "ecommerce"
    },
    "Community App": {
        "template": "Community App Tech Proposal.docx",
        "pricing_fields": [
            ("Design", "design"),
            ("AI/ML & Development", "develop"),
            ("QA & Project Management", "qa_manag"),
            ("Additional Features & Enhancements", "add_fea")
        ],
        "team_fields": [
            ("Project Manager", "pm_no"),
            ("Business Analyst", "ba_no"),
            ("UI/UX Members", "uix_no"),
            ("Backend Developers", "bd_no"),
            ("Frontend Developers", "fd_no"),
            ("AI/ML Developers", "aiml_no"),
            ("System Architect", "sa_no"),
            ("AWS Developer", "aws_no")
        ],
        "special_fields": [
            ("validity_date", "<<"),
            ("location", "<<")
        ],
        "team_type": "community_app"
    },
    "Job Portal Website": {
        "template": "Job portal website Tech Proposal.docx",
        "pricing_fields": [
            ("Design", "design"),
            ("Development", "develop"),
            ("Automations", "autom"),
            ("Testing & Deployment", "deplo"),
            ("Additional Features & Enhancements", "ad_fs")
        ],
        "team_fields": [
            ("Project Manager", "pm_no"),
            ("Business Analyst", "ba_no"),
            ("UI/UX Members", "uix_no"),
            ("Backend Developers", "bd_no"),
            ("Frontend Developers", "fd_no"),
            ("AI/ML Developers", "aiml_no"),
            ("System Architect", "sa_no"),
            ("AWS Developer", "aws_no")
        ],
        "special_fields": [
            ("validity_date", "<<"),
            ("location", "<<")
        ],
        "team_type": "job_portal"
    }
}

def apply_formatting(new_run, original_run):
    """Copy formatting from original run to new run"""
    if original_run.font.name:
        new_run.font.name = original_run.font.name
        new_run._element.rPr.rFonts.set(qn('w:eastAsia'), original_run.font.name)
    if original_run.font.size:
        new_run.font.size = original_run.font.size
    if original_run.font.color.rgb:
        new_run.font.color.rgb = original_run.font.color.rgb
    new_run.bold = original_run.bold
    new_run.italic = original_run.italic

def replace_in_paragraph(para, placeholders):
    """Handle paragraph replacements preserving formatting"""
    original_runs = para.runs.copy()
    full_text = para.text
    for ph, value in placeholders.items():
        full_text = full_text.replace(ph, str(value))

    if full_text != para.text:
        para.clear()
        new_run = para.add_run(full_text)
        if original_runs:
            original_run = next((r for r in original_runs if r.text), None)
            if original_run:
                apply_formatting(new_run, original_run)

def replace_and_format(doc, placeholders):
    """Enhanced replacement with table cell handling"""
    # Process paragraphs
    for para in doc.paragraphs:
        replace_in_paragraph(para, placeholders)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.tables:
                    for nested_table in cell.tables:
                        for nested_row in nested_table.rows:
                            for nested_cell in nested_row.cells:
                                for para in nested_cell.paragraphs:
                                    replace_in_paragraph(para, placeholders)
                else:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para, placeholders)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return doc

def get_marketing_team_details():
    """Collect team composition details specifically for marketing proposals"""
    st.subheader("Marketing Team Composition")
    team_roles = {
        "Project Manager": "PM",
        "Content Writers": "CW",
        "Graphic Designer": "GD",
        "SEO Specialists": "SE",
        "Social Media Manager": "SM",
        "Ad Campaign Manager": "AC"
    }
    team_details = {}
    cols = st.columns(3)

    for idx, (role, placeholder) in enumerate(team_roles.items()):
        with cols[idx % 3]:
            count = st.number_input(
                f"{role} Count:",
                min_value=0,
                step=1,
                key=f"marketing_team_{placeholder}"
            )
            team_details[f"<<{placeholder}>>"] = str(count)
    return team_details

def get_general_team_details():
    """Collect team composition for non-marketing proposals"""
    st.subheader("Team Composition")
    team_roles = {
        "Project Manager": "P1",
        "Frontend Developers": "F1",
        "Business Analyst": "B1",
        "AI/ML Developers": "A1",
        "UI/UX Members": "U1",
        "System Architect": "S1",
        "Backend Developers": "BD1",
        "AWS Developer": "AD1"
    }
    team_details = {}
    cols = st.columns(2)

    for idx, (role, placeholder) in enumerate(team_roles.items()):
        with cols[idx % 2]:
            count = st.number_input(
                f"{role} Count:",
                min_value=0,
                step=1,
                key=f"team_{placeholder}"
            )
            team_details[f"<<{placeholder}>>"] = str(count)
    return team_details

def get_digital_marketing_team_details():
    """Collect team composition details specifically for digital marketing proposals"""
    st.subheader("Digital Marketing Team Composition")
    team_details = {}
    cols = st.columns(2)

    team_roles = {
        "Digital Marketing Executive": "dm_ex_no",
        "Digital Marketing Associate": "dm_asso_no",
        "Business Analyst": "ba_no",
        "Graphics Designer": "gd_no"
    }

    for idx, (role, placeholder) in enumerate(team_roles.items()):
        with cols[idx % 2]:
            count = st.number_input(
                f"{role} Count:",
                min_value=0,
                step=1,
                key=f"dm_team_{placeholder}"
            )
            team_details[f"<<{placeholder}>>"] = str(count)
    return team_details

def get_shopify_team_details():
    """Collect team composition details specifically for Shopify projects"""
    st.subheader("Shopify Team Composition")
    team_details = {}
    cols = st.columns(2)

    team_roles = {
        "Project Manager": "pm_no",
        "Business Analyst": "ba_no",
        "UI/UX Members": "uix_no",
        "Backend Developers": "bd_no",
        "Frontend Developers": "fd_no",
        "AI/ML Developers": "aiml_no",
        "System Architect": "sa_no",
        "Shopify Developers": "sd_no"
    }

    for idx, (role, placeholder) in enumerate(team_roles.items()):
        with cols[idx % 2]:
            count = st.number_input(
                f"{role} Count:",
                min_value=0,
                step=1,
                key=f"shopify_team_{placeholder}"
            )
            team_details[f"<<{placeholder}>>"] = str(count)
    return team_details

def get_fintech_team_details():
    """Collect team composition details specifically for Fintech projects"""
    st.subheader("Fintech Team Composition")
    team_details = {}
    cols = st.columns(2)

    team_roles = {
        "Project Manager": "pm_no",
        "Business Analyst": "ba_no",
        "UI/UX Members": "uix_no",
        "Backend Developers": "bd_no",
        "Frontend Developers": "fd_no",
        "AI/ML Developers": "aiml_no",
        "System Architect": "sa_no",
        "AWS Developer": "aws_no"
    }

    for idx, (role, placeholder) in enumerate(team_roles.items()):
        with cols[idx % 2]:
            count = st.number_input(
                f"{role} Count:",
                min_value=0,
                step=1,
                key=f"fintech_team_{placeholder}"
            )
            team_details[f"<<{placeholder}>>"] = str(count)
    return team_details

def get_search_engine_team_details():
    """Collect team composition details specifically for AI Search Engine projects"""
    st.subheader("AI Search Engine Team Composition")
    team_details = {}
    cols = st.columns(2)

    team_roles = {
        "Project Manager": "pm_no",
        "Business Analyst": "ba_no",
        "UI/UX Members": "uix_no",
        "Backend Developers": "bd_no",
        "Frontend Developers": "fd_no",
        "AI/ML Developers": "aiml_no",
        "System Architect": "sa_no",
        "AWS Developer": "aws_no"
    }

    for idx, (role, placeholder) in enumerate(team_roles.items()):
        with cols[idx % 2]:
            count = st.number_input(
                f"{role} Count:",
                min_value=0,
                step=1,
                key=f"search_engine_team_{placeholder}"
            )
            team_details[f"<<{placeholder}>>"] = str(count)
    return team_details

def get_ecommerce_team_details():
    """Collect team composition details specifically for Ecommerce projects"""
    st.subheader("Ecommerce Team Composition")
    team_details = {}
    cols = st.columns(2)

    team_roles = {
        "Project Manager": "pm_no",
        "Business Analyst": "ba_no",
        "UI/UX Members": "uix_no",
        "Backend Developers": "bd_no",
        "Frontend Developers": "fd_no",
        "AI/ML Developers": "aiml_no",
        "System Architect": "sa_no",
        "AWS Developer": "aws_no"
    }

    for idx, (role, placeholder) in enumerate(team_roles.items()):
        with cols[idx % 2]:
            count = st.number_input(
                f"{role} Count:",
                min_value=0,
                step=1,
                key=f"ecommerce_team_{placeholder}"
            )
            team_details[f"<<{placeholder}>>"] = str(count)
    return team_details

def get_community_app_team_details():
    """Collect team composition details specifically for Community App projects"""
    st.subheader("Community App Team Composition")
    team_details = {}
    cols = st.columns(2)

    team_roles = {
        "Project Manager": "pm_no",
        "Business Analyst": "ba_no",
        "UI/UX Members": "uix_no",
        "Backend Developers": "bd_no",
        "Frontend Developers": "fd_no",
        "AI/ML Developers": "aiml_no",
        "System Architect": "sa_no",
        "AWS Developer": "aws_no"
    }

    for idx, (role, placeholder) in enumerate(team_roles.items()):
        with cols[idx % 2]:
            count = st.number_input(
                f"{role} Count:",
                min_value=0,
                step=1,
                key=f"community_app_team_{placeholder}"
            )
            team_details[f"<<{placeholder}>>"] = str(count)
    return team_details

def get_job_portal_team_details():
    """Collect team composition details specifically for Job Portal projects"""
    st.subheader("Job Portal Team Composition")
    team_details = {}
    cols = st.columns(2)

    team_roles = {
        "Project Manager": "pm_no",
        "Business Analyst": "ba_no",
        "UI/UX Members": "uix_no",
        "Backend Developers": "bd_no",
        "Frontend Developers": "fd_no",
        "AI/ML Developers": "aiml_no",
        "System Architect": "sa_no",
        "AWS Developer": "aws_no"
    }

    for idx, (role, placeholder) in enumerate(team_roles.items()):
        with cols[idx % 2]:
            count = st.number_input(
                f"{role} Count:",
                min_value=0,
                step=1,
                key=f"job_portal_team_{placeholder}"
            )
            team_details[f"<<{placeholder}>>"] = str(count)
    return team_details

def remove_empty_rows(table):
    """Remove rows from the table where the pricing cell is empty or zero"""
    rows_to_remove = []
    for row in table.rows:
        # Skip header row if it exists
        if row.cells[0].text.strip().lower() == 'description':
            continue
            
        # Check if price cell is empty or contains only currency symbol or zero
        if len(row.cells) > 2:  # Ensure we have enough cells
            price_cell = row.cells[2].text.strip()
            if price_cell == "" or price_cell == "$0" or price_cell == "₹0" or price_cell == "0":
                rows_to_remove.append(row)
    
    # Remove rows in reverse order to avoid index issues
    for row in reversed(rows_to_remove):
        table._tbl.remove(row._element)

def validate_phone_number(country, phone_number):
    """Validate phone number based on country"""
    if country.lower() == "india":
        if not phone_number.startswith("+91"):
            return False
    elif country.lower() == "australia":
        if not phone_number.startswith("+61"):
            return False
    else:  # USA and others
        if not phone_number.startswith("+1"):
            return False
    return True

def format_number_with_commas(number):
    """Format number with commas (e.g., 10000 -> 10,000)"""
    return f"{number:,}"

def generate_document():
    st.title("Proposal Generator")
    base_dir = os.getcwd()  # Changed from templates directory

    selected_proposal = st.selectbox("Select Proposal", list(PROPOSAL_CONFIG.keys()))
    config = PROPOSAL_CONFIG[selected_proposal]
    template_path = os.path.join(base_dir, config["template"])

    # Client Information
    col1, col2 = st.columns(2)
    with col1:
        client_name = st.text_input("Client Name:")
        client_email = st.text_input("Client Email:")
    with col2:
        country = st.text_input("Country:")
        client_number = st.text_input("Client Number:")
        if client_number and country:
            if not validate_phone_number(country, client_number):
                prefix = {
                    "india": "+91",
                    "australia": "+61",
                    "usa": "+1"
                }.get(country.lower(), "+1")
                st.error(f"Phone number for {country} should start with {prefix}")

    date_field = st.date_input("Date:", datetime.today())

    # Currency Handling
    currency = st.selectbox("Select Currency", ["USD", "INR", "AUD"])
    currency_symbol = {
        "USD": "$",
        "INR": "₹",
        "AUD": "A$"
    }.get(currency, "$")

    # Special Fields Handling
    special_data = {}
    if config.get("special_fields"):
        st.subheader("Additional Details")
        for field, wrapper in config["special_fields"]:
            if wrapper == "<<":
                placeholder = f"<<{field}>>"
                if field == "VDate" or field == "validity_date":  # Handle both VDate and validity_date
                    vdate = st.date_input(
                        "Proposal Validity Until:",
                        min_value=datetime.today(),
                        value=datetime.today()
                    )
                    special_data[placeholder] = vdate.strftime("%d-%m-%Y")
                elif field in ["advnc_pay", "balnc_pay"]:  # Handle payment fields
                    value = st.number_input(
                        f"{field.replace('_', ' ').title()} ({currency})",
                        min_value=0,
                        step=100,
                        format="%d"
                    )
                    special_data[placeholder] = f"{currency_symbol}{format_number_with_commas(value)}" if value > 0 else ""
                else:
                    special_data[placeholder] = st.text_input(f"{field.replace('_', ' ').title()}:")

    # Initialize placeholders at the start
    placeholders = {
        "<<client_name>>": client_name,
        "<<client_phone>>": client_number,
        "<<client_phoneno>>": client_number,
        "<<client_email>>": client_email,
        "<<date>>": date_field.strftime("%d-%m-%Y"),
        "<<Country>>": country,
        "<<Client Name>>": client_name,
        "<<Client Email>>": client_email,
        "<<Client Number>>": client_number,
        "<<Date>>": date_field.strftime("%d-%m-%Y")
    }

    # Pricing Section
    st.subheader("Pricing Details")
    pricing_data = {}
    numerical_values = {}

    # Create a container for pricing fields
    pricing_container = st.container()
    
    with pricing_container:
        st.write("Enter pricing for each service:")
        
        for label, key in config["pricing_fields"]:
            col1, col2 = st.columns([3, 2])
            with col1:
                st.write(f"• {label}")
            with col2:
                value = st.number_input(
                    f"Amount ({currency})",
                    min_value=0,
                    value=0,
                    step=100,
                    format="%d",
                    key=f"price_{key}"
                )
                numerical_values[key] = value
                if value > 0:
                    pricing_data[f"<<{key}>>"] = f"{currency_symbol}{format_number_with_commas(value)}"
                else:
                    pricing_data[f"<<{key}>>"] = ""

        st.markdown("---")

        # Initialize total at the start
        total = 0

        # Different display for different proposal types
        if selected_proposal == "Web Based AI Fintech":
            # Calculate base total (excluding additional features)
            base_total = (
                numerical_values.get("design", 0) + 
                numerical_values.get("development", 0) + 
                numerical_values.get("ai_ml_model", 0)
            )
            
            # Calculate annual maintenance (10% of base total)
            am_price = int(base_total * 0.10)
            total = base_total + am_price
            
            # Display breakdowns and update pricing data
            st.write(f"**Base Services Cost:** {currency_symbol}{format_number_with_commas(base_total)}")
            st.write(f"**Annual Maintenance (10%):** {currency_symbol}{format_number_with_commas(am_price)}")
            st.write("---")
            
            # Update pricing data for Web Based AI Fintech
            pricing_data["<<annual_main>>"] = f"{currency_symbol}{format_number_with_commas(am_price)}"
            
        elif selected_proposal == "AI Based Search Engine":
            # Calculate base total (excluding additional features)
            base_total = (
                numerical_values.get("design", 0) + 
                numerical_values.get("development", 0) + 
                numerical_values.get("test_deploy", 0)
            )
            
            # Calculate annual maintenance (10% of base total)
            am_price = int(base_total * 0.10)
            total = base_total + am_price
            
            # Display breakdowns
            st.write(f"**Base Services Cost:** {currency_symbol}{format_number_with_commas(base_total)}")
            st.write(f"**Annual Maintenance (10%):** {currency_symbol}{format_number_with_commas(am_price)}")
            st.write("---")
            
            # Update pricing data for AI Based Search Engine
            pricing_data["<<annual_mainte>>"] = f"{currency_symbol}{format_number_with_commas(am_price)}"
            
        elif selected_proposal == "Single Vendor Ecommerce":
            # Calculate base total (excluding additional features)
            base_total = (
                numerical_values.get("design", 0) + 
                numerical_values.get("dev", 0) + 
                numerical_values.get("wb_cb", 0) + 
                numerical_values.get("testing", 0)
            )
            
            # Calculate annual maintenance (10% of base total)
            am_price = int(base_total * 0.10)
            
            # Calculate final total
            total = base_total + am_price
            
            # Display breakdowns
            st.write(f"**Base Services Cost:** {currency_symbol}{format_number_with_commas(base_total)}")
            st.write(f"**Annual Maintenance (10%):** {currency_symbol}{format_number_with_commas(am_price)}")
            st.write("---")
            if currency == "INR":
                st.write(f"**Total Amount:** {currency_symbol}{format_number_with_commas(total)} + 18% GST")
            else:
                st.write(f"**Total Amount:** {currency_symbol}{format_number_with_commas(total)}")
            
            # Additional Features display (separate from total)
            af_price = {
                "USD": 250,
                "INR": 25000,
                "AUD": 375
            }.get(currency, 250)
            st.write("---")
            st.write(f"**Additional Features & Enhancements:** {currency_symbol}{format_number_with_commas(af_price)}")
            
            # Update pricing data
            pricing_data["<<an_ma>>"] = f"{currency_symbol}{format_number_with_commas(am_price)}"
            total_price_str = f"{currency_symbol}{format_number_with_commas(total)}"
            if currency == "INR":
                total_price_str += " + 18% GST"
            pricing_data["<<total>>"] = total_price_str
            placeholders["<<total>>"] = total_price_str
            pricing_data["<<ad_fs>>"] = f"{currency_symbol}{format_number_with_commas(af_price)}"

        elif selected_proposal == "Community App":
            # Calculate base total (excluding additional features)
            base_total = (
                numerical_values.get("design", 0) + 
                numerical_values.get("develop", 0) + 
                numerical_values.get("qa_manag", 0)
            )
            
            # Calculate annual maintenance (10% of base total)
            am_price = int(base_total * 0.10)
            
            # Calculate final total
            total = base_total + am_price
            
            # Display breakdowns
            st.write(f"**Base Services Cost:** {currency_symbol}{format_number_with_commas(base_total)}")
            st.write(f"**Annual Maintenance (10%):** {currency_symbol}{format_number_with_commas(am_price)}")
            st.write("---")
            
            # Format total price
            total_price_str = f"{currency_symbol}{format_number_with_commas(total)}"
            if currency == "INR":
                total_price_str += " + 18% GST"
            st.write(f"**Total Amount:** {total_price_str}")
            
            # Additional Features price
            af_price = {
                "USD": 250,
                "INR": 25000,
                "AUD": 375
            }.get(currency, 250)
            st.write("---")
            st.write(f"**Additional Features & Enhancements:** {currency_symbol}{format_number_with_commas(af_price)}")
            
            # Update pricing data
            pricing_data.update({
                "<<design>>": f"{currency_symbol}{format_number_with_commas(numerical_values.get('design', 0))}",
                "<<develop>>": f"{currency_symbol}{format_number_with_commas(numerical_values.get('develop', 0))}",
                "<<qa_manag>>": f"{currency_symbol}{format_number_with_commas(numerical_values.get('qa_manag', 0))}",
                "<<ann_main>>": f"{currency_symbol}{format_number_with_commas(am_price)}",
                "<<total_price>>": total_price_str,
                "<<add_fea>>": f"{currency_symbol}{format_number_with_commas(af_price)}",
                "<<AF-Price>>": f"{currency_symbol}{format_number_with_commas(af_price)}"
            })
            
            # Update placeholders
            placeholders.update(pricing_data)
            
        elif selected_proposal == "Job Portal Website":
            # Calculate base total (excluding additional features)
            base_total = (
                numerical_values.get("design", 0) + 
                numerical_values.get("develop", 0) + 
                numerical_values.get("autom", 0) + 
                numerical_values.get("deplo", 0)
            )
            
            # Calculate annual maintenance (10% of base total)
            am_price = int(base_total * 0.10)
            
            # Calculate final total
            total = base_total + am_price
            
            # Display breakdowns
            st.write(f"**Base Services Cost:** {currency_symbol}{format_number_with_commas(base_total)}")
            st.write(f"**Annual Maintenance (10%):** {currency_symbol}{format_number_with_commas(am_price)}")
            st.write("---")
            if currency == "INR":
                st.write(f"**Total Amount:** {currency_symbol}{format_number_with_commas(total)} + 18% GST")
            else:
                st.write(f"**Total Amount:** {currency_symbol}{format_number_with_commas(total)}")
            
            # Additional Features display (separate from total)
            af_price = {
                "USD": 250,
                "INR": 25000,
                "AUD": 375
            }.get(currency, 250)
            st.write("---")
            st.write(f"**Additional Features & Enhancements:** {currency_symbol}{format_number_with_commas(af_price)}")
            
            # Update pricing data
            pricing_data.update({
                "<<design>>": f"{currency_symbol}{format_number_with_commas(numerical_values.get('design', 0))}",
                "<<develop>>": f"{currency_symbol}{format_number_with_commas(numerical_values.get('develop', 0))}",
                "<<autom>>": f"{currency_symbol}{format_number_with_commas(numerical_values.get('autom', 0))}",
                "<<deplo>>": f"{currency_symbol}{format_number_with_commas(numerical_values.get('deplo', 0))}",
                "<<an_m>>": f"{currency_symbol}{format_number_with_commas(am_price)}",
                "<<total_price>>": f"{currency_symbol}{format_number_with_commas(total)}" + (" + 18% GST" if currency == "INR" else ""),
                "<<ad_fs>>": f"{currency_symbol}{format_number_with_commas(af_price)}"
            })
            
            # Update placeholders
            placeholders.update(pricing_data)
            
        else:
            # Original calculation for other proposals
            total = sum(numerical_values.values())
            am_price = int(total * 0.10)
            total += am_price
            
            # Additional Features price
            af_price = {
                "USD": 250,
                "INR": 25000,
                "AUD": 375
            }.get(currency, 250)
            
            # Update pricing data
            pricing_data["<<AM-Price>>"] = f"{currency_symbol}{format_number_with_commas(am_price)}"
            pricing_data["<<AF-Price>>"] = f"{currency_symbol}{format_number_with_commas(af_price)}"
            
            if currency == "INR":
                pricing_data["<<T-Price>>"] = f"{currency_symbol}{format_number_with_commas(total)} + 18% GST"
            else:
                pricing_data["<<T-Price>>"] = f"{currency_symbol}{format_number_with_commas(total)}"

    # Team Composition
    team_data = {}
    if config["team_type"] == "marketing":
        team_data = get_marketing_team_details()
    elif config["team_type"] == "general":
        team_data = get_general_team_details()
    elif config["team_type"] == "digital_marketing":
        team_data = get_digital_marketing_team_details()
    elif config["team_type"] == "shopify":
        team_data = get_shopify_team_details()
    elif config["team_type"] == "fintech":
        team_data = get_fintech_team_details()
    elif config["team_type"] == "search_engine":
        team_data = get_search_engine_team_details()
    elif config["team_type"] == "ecommerce":
        team_data = get_ecommerce_team_details()
    elif config["team_type"] == "community_app":
        team_data = get_community_app_team_details()
    elif config["team_type"] == "job_portal":
        team_data = get_job_portal_team_details()

    # Add Additional Tools Section
    st.subheader("Add Additional Tools")
    additional_tool_1 = st.text_input("Tool 1:")
    additional_tool_2 = st.text_input("Tool 2:")

    additional_tools_data = {}
    if additional_tool_1:
        additional_tools_data["<<T1>>"] = additional_tool_1
    else:
        additional_tools_data["<<T1>>"] = ""

    if additional_tool_2:
        additional_tools_data["<<T2>>"] = additional_tool_2
    else:
        additional_tools_data["<<T2>>"] = ""

    # Calculate total price string based on currency
    total_price_str = f"{currency_symbol}{format_number_with_commas(total)}"
    if currency == "INR":
        total_price_str += " + 18% GST"

    # Update all placeholders
    placeholders.update(pricing_data)
    placeholders.update(team_data)
    placeholders.update(special_data)
    placeholders.update(additional_tools_data)

    if st.button("Generate Proposal"):
        if client_number and country and not validate_phone_number(country, client_number):
            st.error(f"Invalid phone number format for {country} should start with {'+91' if country.lower() == 'india' else '+1'}.")
        else:
            formatted_date = date_field.strftime("%d %b %Y")
            unique_id = str(uuid.uuid4())[:8]
            doc_filename = f"{selected_proposal}_{client_name}_{formatted_date}_{unique_id}.docx"

            with tempfile.TemporaryDirectory() as temp_dir:
                try:
                    doc = Document(template_path)
                except FileNotFoundError:
                    st.error(f"Template file not found: {template_path}")
                    return

                doc = replace_and_format(doc, placeholders)

                # Remove empty rows from the pricing table
                for table in doc.tables:
                    remove_empty_rows(table)

                doc_path = os.path.join(temp_dir, doc_filename)
                doc.save(doc_path)

                with open(doc_path, "rb") as f:
                    st.download_button(
                        label="Download Proposal",
                        data=f,
                        file_name=doc_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

if __name__ == "__main__":
    generate_document()
